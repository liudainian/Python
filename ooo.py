#coding:utf-8

import xlrd
import docx
import docx.styles
from docx.shared import RGBColor
import MySQLdb
from datetime import date

class ExcelMethod:
    '''excel basic class'''
    count= 0 #全局变量
    def __init__(self,path):
        self.path = path
        self.workbook = None

    def openExcel(self):
        self.workbook = xlrd.open_workbook(self.path)

    def getSheelObject(self,sheetindex,sheetname):
        sheetlist = self.workbook.sheet_names()#取得sheet列表
        if sheetname in sheetlist:
            sheet_object = self.workbook.sheet_by_name(sheetname) #按表名获取sheel表中的内容
            print (sheet_object)
            return sheet_object
        else:
            sheet_object = self.workbook.sheet_by_index(sheetindex)#按索引获取sheel表中的内容
            sheet_object.ncols
            return sheet_object

    def operationSheet(self, sheetObject, cellRowNum , cellColNum):
        # sheet的名称，行数，列数column number
        sheetname = sheetObject.name
        sheetrows = sheetObject.nrows
        sheetcols = sheetObject.ncols
        # 获取单元格内容的数据类型 ctype : 0 empty,1 string, 2 number, 3 date, 4 boolean, 5 error
        celltype = sheetObject.cell(cellRowNum, cellColNum).ctype
        #获取单元格内容
        if celltype == 1:
            cellvalue = sheetObject.cell_value(cellRowNum, cellColNum).decode("utf-8")
        elif celltype == 3:
            cellvalue_tuple = xlrd.xldate_as_tuple(sheetObject.cell_value(cellRowNum, cellColNum),self.workbook.datemode)#返回值：(2018, 5, 25, 0, 0, 0)
            cellvalue = date(*cellvalue_tuple[:3]).strftime('%Y/%m/%d')#取前面三位  以 / 连接时间
        else:
            cellvalue = sheetObject.cell_value(cellRowNum,cellColNum)
        print celltype
        return cellvalue

class WordMethod:
    def __init__(self, path):
        self.path = path
        self.docx = None

    def openWord(self):
        self.docx = docx.Document(self.path)

    def readFromWord(self):
        content = '\n'.join([para.text for para in self.docx.paragraphs])
        return content

    def writeToWord(self):

        # 新增样式(第一个参数是样式名称，第二个参数是样式类型：1代表段落；2代表字符；3代表表格)
        style = self.docx.styles.add_style('style name 1', 2)

        # 设置具体样式(修改样式字体为蓝色，当然还可以修改其他的，大家自己尝试)
        style.font.color.rgb = RGBColor(0xff, 0x0, 0x0)

        self.docx.add_paragraph(u'第一段', style=None)  # 插入一个段落，文本为“第一段”

    def saveToNewWord(self, saveAs = ''):
        if saveAs == "": #没有输入newpaath则直接替换当前文件
            saveAs = self.path
        else:     #另存文件到newpath
            pass
        self.docx.save(saveAs)

if __name__ == "__main__":
    path = "C:\Users\hp\Desktop\\test.xlsx"
    excel = ExcelMethod(path)
    excel.openExcel()
    workobject = excel.getSheelObject(1,'Sheet1')
    cellvalue = excel.operationSheet(workobject,2,5)
    print cellvalue

    path = "C:\Users\hp\Desktop\关于xx销售大区xxx窜货处理通报OA.docx".decode('utf-8')
    word = WordMethod(path)
    word.openWord()
    value = word.readFromWord()
    word.writeToWord()
    word.saveToNewWord("C:\Users\hp\Desktop\\1.docx")
    print value
