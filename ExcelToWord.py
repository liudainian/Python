#coding:utf-8

import xlrd
import xlwt
import docx
from datetime import date

def readFromExcel(path):
    workbook = xlrd.open_workbook(path)
    print(workbook.sheet_names())  #row number
    # 根据sheet索引或者名称获取sheet内容
    sheet2 = workbook.sheet_by_index(1)  # sheet索引从0开始
    sheet2 = workbook.sheet_by_name('Sheet2')

    # sheet的名称，行数，列数column number
    print (sheet2.name, sheet2.nrows, sheet2.ncols)

    # 获取整行和整列的值（数组）
    rows = sheet2.row_values(2)  # 获取第三行内容
    cols = sheet2.col_values(2)  # 获取第三列内容
    print "rows" ,rows
    print "cols",cols

    # 获取单元格内容
    print sheet2.cell(2, 0).value.encode('utf-8')
    print sheet2.cell_value(2, 1).encode('utf-8')
    print sheet2.row(2)[0].value.encode('utf-8')

    # 获取单元格内容的数据类型 ctype : 0 empty,1 string, 2 number, 3 date, 4 boolean, 5 error
    print sheet2.cell(2, 0).ctype

    #处理日期显示问题 2018/2/22
    row = 2;col = 24
    if (sheet2.cell(row, col).ctype == 3): #row 行号-1
        date_value = xlrd.xldate_as_tuple(sheet2.cell_value(rows, 3), workbook.datemode)
        date_tmp = date(*date_value[:3]).strftime('%Y/%m/%d')

    #处理合并单元格

    #workbook = xlrd.open_workbook(path, formatting_info=True)#参数为true 返回索引号  支持版本excel2005-2012
    #sheet2 = workbook.sheet_by_name('Sheet2')
    # 输出参数为merged_cells返回的这四个参数的含义是：(row,row_range,col,col_range),
    #其中[row,row_range)包括row,不包括row_range,
    print sheet2.merged_cells
    #获取合并单元格的索引
    merge = []
    for (rlow, rhigh, clow, chigh) in sheet2.merged_cells:
        merge.append([rlow, clow])
    print(merge)

def readFromWord(path):
    document = docx.Document(path)
    docText = '\n\n'.join([
        paragraph.text.encode('utf-8') for paragraph in document.paragraphs
    ])
    print(docText)

def writeToWord(path):
    from docx import Document
    from docx.shared import Inches

    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )

    #document.add_picture('monty-truth.png', width=Inches(1.25))

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    '''for item in recordset:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.qty)
        row_cells[1].text = str(item.id)
        row_cells[2].text = item.desc'''

    document.add_page_break()

    document.save(path)


if __name__ == "__main__":
    path = "C:\Users\Administrator\Desktop\\test.xlsx"
    #readFromExcel(path)
    path = "C:\Users\Administrator\Desktop\\RobotFramework问题解答.docx".decode("utf-8")
    #readFromWord(path)
    path = 'C:\Users\Administrator\Desktop\\ddd.docx'
    writeToWord(path)